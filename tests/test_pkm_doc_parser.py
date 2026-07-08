"""Unit tests for apps.pkm.services.doc_parser.

These tests create small sample files (PDF, DOCX, TXT, MD, XLSX) in the
pytest-managed ``tmp_path`` fixture, then verify that
:func:`extract_text` returns the expected content for each supported type and
raises :class:`ValueError` for unsupported types.

No database access is required (pure parsing logic), but we still configure
Django settings via ``DJANGO_SETTINGS_MODULE`` for consistency with the rest
of the test suite.  Each test is independent and cleans up via ``tmp_path``.
"""

from __future__ import annotations

import pytest

from apps.pkm.services import doc_parser
from apps.pkm.services.doc_parser import extract_text

# ---------------------------------------------------------------------------
# Helpers to build sample fixture files
# ---------------------------------------------------------------------------


def _make_pdf(path, text: str) -> None:
    """Create a minimal one-page PDF containing ``text`` using pypdf.

    We add a standard Helvetica Type1 font to the page resources and inject a
    text-showing operator (``BT ... Tj ET``) into the content stream.  This
    produces a real, parseable PDF whose extracted text round-trips through
    ``PdfReader.extract_text()``.
    """
    from pypdf import PdfWriter
    from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)

    # Register a built-in Type1 font (Helvetica) so the Tj operator can show text.
    font_dict = DictionaryObject()
    font_dict[NameObject("/Type")] = NameObject("/Font")
    font_dict[NameObject("/Subtype")] = NameObject("/Type1")
    font_dict[NameObject("/BaseFont")] = NameObject("/Helvetica")
    font_ref = writer._add_object(font_dict)

    resources = page["/Resources"]
    resources[NameObject("/Font")] = DictionaryObject()
    resources["/Font"][NameObject("/F1")] = font_ref

    # Escape parentheses and backslashes for the PDF string literal.
    escaped = text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
    stream_data = f"BT /F1 12 Tf 72 700 Td ({escaped}) Tj ET".encode("latin-1", errors="replace")
    cs = DecodedStreamObject()
    cs.set_data(stream_data)
    page[NameObject("/Contents")] = writer._add_object(cs)

    with open(path, "wb") as f:
        writer.write(f)


def _make_docx(path, text: str) -> None:
    """Create a small .docx whose first paragraph is ``text``."""
    import docx

    document = docx.Document()
    document.add_paragraph(text)
    document.save(str(path))


def _make_txt(path, text: str) -> None:
    path.write_text(text, encoding="utf-8")


def _make_xlsx(path, rows: list[list[str]]) -> None:
    """Create a small .xlsx with the given cell values starting at A1."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    for row in rows:
        ws.append(row)
    wb.save(str(path))
    wb.close()


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------


def test_extract_pdf(tmp_path):
    pdf_path = tmp_path / "sample.pdf"
    _make_pdf(pdf_path, "Hello PKM PDF")
    result = extract_text(str(pdf_path))
    assert "Hello PKM PDF" in result
    assert len(result.strip()) > 0


def test_extract_pdf_multiline(tmp_path):
    pdf_path = tmp_path / "multi.pdf"
    _make_pdf(pdf_path, "Invoice total 1234")
    result = extract_text(str(pdf_path))
    assert "Invoice" in result
    assert "1234" in result


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------


def test_extract_docx(tmp_path):
    docx_path = tmp_path / "sample.docx"
    _make_docx(docx_path, "Knowledge note about accounting.")
    result = extract_text(str(docx_path))
    assert "Knowledge note about accounting." in result
    assert len(result.strip()) > 0


def test_extract_docx_multiple_paragraphs(tmp_path):
    import docx

    path = tmp_path / "multi.docx"
    document = docx.Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    document.add_paragraph("Third paragraph.")
    document.save(str(path))

    result = extract_text(str(path))
    assert "First paragraph." in result
    assert "Second paragraph." in result
    assert "Third paragraph." in result


def test_extract_docx_empty_paragraphs_skipped(tmp_path):
    import docx

    path = tmp_path / "blanks.docx"
    document = docx.Document()
    document.add_paragraph("Real content")
    document.add_paragraph("")  # blank
    document.add_paragraph("   ")  # whitespace only
    document.save(str(path))

    result = extract_text(str(path))
    assert "Real content" in result


# ---------------------------------------------------------------------------
# TXT extraction
# ---------------------------------------------------------------------------


def test_extract_txt(tmp_path):
    txt_path = tmp_path / "sample.txt"
    _make_txt(txt_path, "Plain text content for PKM.")
    result = extract_text(str(txt_path))
    assert result == "Plain text content for PKM."


def test_extract_txt_unicode(tmp_path):
    txt_path = tmp_path / "unicode.txt"
    content = "Tiêu đề ghi chú -测试 - 🔑"
    _make_txt(txt_path, content)
    result = extract_text(str(txt_path))
    assert "Tiêu đề" in result
    assert "测试" in result


def test_extract_txt_bom(tmp_path):
    """A UTF-8 file with BOM should decode cleanly (utf-8-sig fallback)."""
    txt_path = tmp_path / "bom.txt"
    txt_path.write_bytes(b"\xef\xbb\xbfHello BOM world")
    result = extract_text(str(txt_path))
    assert "Hello BOM world" in result


def test_extract_txt_latin1_fallback(tmp_path):
    """Non-UTF-8 bytes should still decode via latin-1 fallback."""
    txt_path = tmp_path / "latin1.txt"
    txt_path.write_bytes(b"Caf\xe9 m\xfcnchen")
    result = extract_text(str(txt_path))
    assert len(result) > 0
    assert "Caf" in result


def test_extract_txt_empty(tmp_path):
    txt_path = tmp_path / "empty.txt"
    _make_txt(txt_path, "")
    assert extract_text(str(txt_path)) == ""


# ---------------------------------------------------------------------------
# Markdown extraction
# ---------------------------------------------------------------------------


def test_extract_md(tmp_path):
    md_path = tmp_path / "note.md"
    content = "# Title\n\nSome **markdown** content."
    _make_txt(md_path, content)
    result = extract_text(str(md_path))
    assert "# Title" in result
    assert "markdown" in result
    assert len(result.strip()) > 0


def test_extract_md_markdown_syntax_preserved(tmp_path):
    """Markdown is read as-is; no rendering should occur."""
    md_path = tmp_path / "syntax.md"
    content = "- item 1\n- item 2\n\n```python\nx = 1\n```"
    _make_txt(md_path, content)
    result = extract_text(str(md_path))
    assert result == content


# ---------------------------------------------------------------------------
# XLSX extraction
# ---------------------------------------------------------------------------


def test_extract_xlsx(tmp_path):
    xlsx_path = tmp_path / "data.xlsx"
    _make_xlsx(
        xlsx_path,
        [["Name", "Amount"], ["Invoice", "1500000"], ["Tax", "150000"]],
    )
    result = extract_text(str(xlsx_path))
    assert "Name" in result
    assert "Amount" in result
    assert "Invoice" in result
    assert "1500000" in result
    assert len(result.strip()) > 0


def test_extract_xlsx_empty_cells_skipped(tmp_path):
    xlsx_path = tmp_path / "blanks.xlsx"
    _make_xlsx(xlsx_path, [["A", None, "C"], [None, None, None]])
    result = extract_text(str(xlsx_path))
    assert "A" in result
    assert "C" in result


def test_extract_xlsx_multiple_sheets(tmp_path):
    from openpyxl import Workbook

    path = tmp_path / "multi_sheet.xlsx"
    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Sheet1"
    ws1.append(["Alpha"])
    ws2 = wb.create_sheet("Sheet2")
    ws2.append(["Beta"])
    wb.save(str(path))
    wb.close()

    result = extract_text(str(path))
    assert "Alpha" in result
    assert "Beta" in result


# ---------------------------------------------------------------------------
# Unsupported / error cases
# ---------------------------------------------------------------------------


def test_unsupported_extension_raises_value_error(tmp_path):
    fake_path = tmp_path / "archive.zip"
    fake_path.write_bytes(b"PK\x03\x04 fake zip")
    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text(str(fake_path))


def test_no_extension_raises_value_error(tmp_path):
    no_ext = tmp_path / "README"
    no_ext.write_text("no extension")
    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text(str(no_ext))


def test_unsupported_exe_raises_value_error(tmp_path):
    exe_path = tmp_path / "malware.exe"
    exe_path.write_bytes(b"MZ fake exe")
    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text(str(exe_path))


def test_missing_file_raises_filenotfound(tmp_path):
    missing = tmp_path / "does_not_exist.pdf"
    with pytest.raises(FileNotFoundError):
        extract_text(str(missing))


def test_extension_case_insensitive(tmp_path):
    """Uppercase extensions should be accepted."""
    txt_path = tmp_path / "upper.TXT"
    _make_txt(txt_path, "Uppercase ext")
    assert extract_text(str(txt_path)) == "Uppercase ext"


def test_mixed_case_extension(tmp_path):
    md_path = tmp_path / "mixed.Md"
    _make_txt(md_path, "Mixed case markdown")
    result = extract_text(str(md_path))
    assert "Mixed case markdown" in result


# ---------------------------------------------------------------------------
# Supported extensions constant
# ---------------------------------------------------------------------------


def test_supported_extensions_contains_all():
    expected = {"pdf", "docx", "txt", "md", "xlsx"}
    assert set(doc_parser.SUPPORTED_EXTENSIONS) == expected
