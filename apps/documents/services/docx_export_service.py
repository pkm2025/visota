"""DOCX export service — generate Word documents from voucher/contract data."""

from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


class DocxExportService:
    """Generate .docx files for vouchers, contracts, reports."""

    def export_voucher(self, voucher) -> bytes:
        """Export an accounting voucher as DOCX."""
        doc = Document()

        # Page setup
        section = doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        # Company header
        company = voucher.company
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(company.name)
        run.bold = True
        run.font.size = Pt(14)

        if company.address:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(f"Địa chỉ: {company.address}")
            run2.font.size = Pt(10)
        if company.tax_code:
            p3 = doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run3 = p3.add_run(f"MST: {company.tax_code}")
            run3.font.size = Pt(10)

        # Title
        doc.add_paragraph()
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(voucher.get_voucher_type_display())
        title_run.bold = True
        title_run.font.size = Pt(16)

        if voucher.description:
            desc = doc.add_paragraph()
            desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desc.add_run(voucher.description).font.size = Pt(12)

        # Voucher info table
        doc.add_paragraph()
        info_table = doc.add_table(rows=2, cols=4)
        info_table.style = "Table Grid"

        cells = info_table.rows[0].cells
        cells[0].text = "Số chứng từ"
        cells[1].text = voucher.voucher_no
        cells[2].text = "Ngày"
        cells[3].text = voucher.voucher_date.strftime("%d/%m/%Y")

        cells = info_table.rows[1].cells
        cells[2].text = "Trạng thái"
        cells[3].text = voucher.get_status_display()

        # Lines table
        doc.add_paragraph()
        doc.add_paragraph("Bút toán (hạch toán):").bold = True

        lines = list(voucher.lines.all())
        table = doc.add_table(rows=1 + len(lines) + 1, cols=6)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        headers = ["#", "TK Nợ", "TK Có", "Đối tượng", "Diễn giải", "Số tiền"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows
        for idx, line in enumerate(lines, 1):
            row = table.rows[idx]
            row.cells[0].text = str(line.line_no)
            row.cells[1].text = line.account_code if line.debit_vnd > 0 else ""
            row.cells[2].text = line.account_code if line.credit_vnd > 0 else ""
            row.cells[3].text = line.object_code or ""
            row.cells[4].text = line.description or ""
            amount = line.debit_vnd or line.credit_vnd
            row.cells[5].text = f"{amount:,.0f}"

        # Total row
        total_row = table.rows[-1]
        total_row.cells[0].text = ""
        total_row.cells[4].text = "Tổng cộng"
        total_row.cells[5].text = f"{voucher.total_vnd:,.0f}"
        for cell in total_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Amount in words
        doc.add_paragraph()
        doc.add_paragraph(
            "Số tiền viết bằng chữ: ......................................................."
        ).italic = True

        # Signatures
        doc.add_paragraph()
        sig_table = doc.add_table(rows=1, cols=3)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        labels = [
            "Người lập\n(Ký, họ tên)",
            "Kế toán trưởng\n(Ký, họ tên)",
            "Giám đốc\n(Ký, họ tên, đóng dấu)",
        ]
        for i, label in enumerate(labels):
            sig_table.rows[0].cells[i].text = label
            for paragraph in sig_table.rows[0].cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def export_contract_from_template(self, contract, template) -> bytes:
        """Export a contract using a ContractTemplate as DOCX.

        Since templates are HTML, this method creates a simplified DOCX
        with the key contract fields in a structured format.
        """
        doc = Document()

        # Company header
        company = contract.company
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(company.name)
        run.bold = True
        run.font.size = Pt(14)

        # Title
        doc.add_paragraph()
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(template.name)
        title_run.bold = True
        title_run.font.size = Pt(16)

        # Contract info
        doc.add_paragraph()
        fields = [
            ("Số hợp đồng", contract.contract_no),
            (
                "Ngày hợp đồng",
                contract.contract_date.strftime("%d/%m/%Y") if contract.contract_date else "",
            ),
            ("Bên A", company.name),
            ("MST", company.tax_code or ""),
            ("Bên B", contract.party_name or ""),
            ("MST bên B", contract.party_tax_code or ""),
            ("Giá trị", f"{contract.value:,.0f} VND" if contract.value else ""),
            (
                "Ngày bắt đầu",
                contract.start_date.strftime("%d/%m/%Y") if contract.start_date else "",
            ),
            (
                "Ngày kết thúc",
                contract.end_date.strftime("%d/%m/%Y") if contract.end_date else "",
            ),
        ]

        for label, value in fields:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            p.add_run(str(value))

        # Legal basis
        if getattr(template, "legal_basis", ""):
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(f"Căn cứ pháp lý: {template.legal_basis}")
            run.italic = True
            run.font.size = Pt(9)

        # Signatures
        doc.add_paragraph()
        doc.add_paragraph()
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sig_table.rows[0].cells[0].text = "BÊN A\n\n\n(Ký, ghi rõ họ tên, đóng dấu)"
        sig_table.rows[0].cells[1].text = "BÊN B\n\n\n(Ký, ghi rõ họ tên, đóng dấu)"
        for cell in sig_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def export_trial_balance(self, balances, fiscal_year, period, totals) -> bytes:
        """Export trial balance as DOCX."""
        doc = Document()

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("BẢNG CÂN ĐỐI TÀI KHOẢN")
        run.bold = True
        run.font.size = Pt(14)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(f"Tháng {period}/{fiscal_year}").font.size = Pt(11)

        # Table
        rows_count = len(balances) + 2  # header + data + totals
        table = doc.add_table(rows=rows_count, cols=7)
        table.style = "Table Grid"

        headers = [
            "TK",
            "SD Đầu Nợ",
            "SD Đầu Có",
            "PS Nợ",
            "PS Có",
            "SD Cuối Nợ",
            "SD Cuối Có",
        ]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for idx, b in enumerate(balances, 1):
            row = table.rows[idx]
            row.cells[0].text = b.account_code
            row.cells[1].text = f"{b.opening_debit or 0:,.0f}"
            row.cells[2].text = f"{b.opening_credit or 0:,.0f}"
            row.cells[3].text = f"{b.period_debit or 0:,.0f}"
            row.cells[4].text = f"{b.period_credit or 0:,.0f}"
            row.cells[5].text = f"{b.closing_debit or 0:,.0f}"
            row.cells[6].text = f"{b.closing_credit or 0:,.0f}"

        # Totals
        total_row = table.rows[-1]
        total_row.cells[0].text = "Tổng"
        total_row.cells[1].text = f"{totals.get('opening_debit', 0):,.0f}"
        total_row.cells[2].text = f"{totals.get('opening_credit', 0):,.0f}"
        total_row.cells[3].text = f"{totals.get('period_debit', 0):,.0f}"
        total_row.cells[4].text = f"{totals.get('period_credit', 0):,.0f}"
        total_row.cells[5].text = f"{totals.get('closing_debit', 0):,.0f}"
        total_row.cells[6].text = f"{totals.get('closing_credit', 0):,.0f}"
        for cell in total_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
